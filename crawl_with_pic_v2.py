
import requests
from readability.readability import Document
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.shared import Inches
import pandas as pd
from io import BytesIO
from PIL import Image
import os
from urllib.parse import urljoin
import re

def sanitize_filename(name):
    return re.sub(r'[\\/*?:"<>|]', "_", name)

def fetch_microsoft_docs_to_word(excel_file, sheet_name, url_column, output_folder):
    # Read URLs from Excel
    df = pd.read_excel(excel_file, sheet_name=sheet_name, engine='openpyxl')
    urls = df[url_column].dropna().tolist()

    # Create output folder
    os.makedirs(output_folder, exist_ok=True)

    for idx, url in enumerate(urls, start=1):
        try:
            print(f"Processing {idx}/{len(urls)}: {url}")
            response = requests.get(url)
            response.raise_for_status()

            # Extract main content using readability-lxml
            readable_doc = Document(response.text)
            cleaned_html = readable_doc.summary()
            page_title = readable_doc.title() or f"Page_{idx}"

            # Parse cleaned HTML
            soup = BeautifulSoup(cleaned_html, 'html.parser')

            # Prepare Word document
            doc = DocxDocument()
            doc.add_heading(page_title, level=0)
            doc.add_paragraph(f"Source URL: {url}")

            # Output file path
            file_name = sanitize_filename(page_title[:100]) + ".docx"
            file_path = os.path.join(output_folder, file_name)

            # Extract elements in order: headings, paragraphs, images
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'p', 'img']):
                if element.name == 'img':
                    img_url = element.get('src') or element.get('data-src')
                    if img_url:
                        img_url = urljoin(url, img_url)
                        try:
                            img_response = requests.get(img_url)
                            img_response.raise_for_status()
                            image = Image.open(BytesIO(img_response.content))
                            image.thumbnail((600, 600))
                            img_stream = BytesIO()
                            image.save(img_stream, format='PNG')
                            img_stream.seek(0)
                            doc.add_picture(img_stream, width=Inches(3))
                        except Exception as img_err:
                            doc.add_paragraph(f"Image could not be added: {img_err}")
                else:
                    text = element.get_text(strip=False)
                    if text:
                        if element.name.startswith('h'):
                            level = int(element.name[1]) if element.name[1].isdigit() else 1
                            doc.add_heading(text, level=level)
                        else:
                            doc.add_paragraph(text)

            # Save Word file
            doc.save(file_path)
            print(f"Saved: {file_path}")

        except Exception as e:
            print(f"Error processing {url}: {e}")

    print(f"All files saved in folder: {output_folder}")

# Example usage
if __name__ == "__main__":
    excel_file = "urls.xlsx"       # Path to your Excel file
    sheet_name = "Sheet1"          # Sheet name
    url_column = "URL"             # Column name containing URLs
    output_folder = "microsoft_docs_pages"

    fetch_microsoft_docs_to_word(excel_file, sheet_name, url_column, output_folder)
