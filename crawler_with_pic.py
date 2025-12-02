
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import pandas as pd
from io import BytesIO
from PIL import Image
import re
import os
from urllib.parse import urljoin
import hashlib

def sanitize_filename(name):
    return re.sub(r'[\\/*?:"<>|]', "_", name)

def text_hash(text):
    return hashlib.md5(text.encode('utf-8')).hexdigest()

def urls_from_excel_div_content(excel_file, sheet_name, url_column, output_folder):
    df = pd.read_excel(excel_file, sheet_name=sheet_name)
    urls = df[url_column].dropna().tolist()

    os.makedirs(output_folder, exist_ok=True)

    for idx, url in enumerate(urls, start=1):
        try:
            response = requests.get(url)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')

            # Remove header, footer, nav, aside
            for tag in soup.find_all(['header', 'footer', 'nav', 'aside']):
                tag.decompose()

            # Get page title for filename
            title_tag = soup.find('title')
            page_title = title_tag.get_text(strip=True) if title_tag else f"Page_{idx}"
            file_name = sanitize_filename(page_title[:100]) + ".docx"
            file_path = os.path.join(output_folder, file_name)

            # Find main content block
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_=re.compile(r'(content|article|post|body)'))
            if not main_content:
                main_content = soup

            doc = Document()
            doc.add_heading(page_title, level=0)
            doc.add_paragraph(f"Source URL: {url}")

            seen_hashes = set()
            for element in main_content.find_all(['div', 'img']):  # Only div and img
                if element.name == 'img':
                    img_url = element.get('src') or element.get('data-src')
                    if img_url:
                        img_url = urljoin(url, img_url)
                        try:
                            img_response = requests.get(img_url)
                            img_response.raise_for_status()
                            image = Image.open(BytesIO(img_response.content))
                            image.thumbnail((600, 600))
                            img_stream = BytesIO()
                            image.save(img_stream, format='PNG')
                            img_stream.seek(0)
                            doc.add_picture(img_stream, width=Inches(3))
                        except Exception:
                            doc.add_paragraph("Image could not be added.")
                else:
                    text = element.get_text(strip=True)
                    if text and len(text) > 30:  # Skip short fragments
                        h = text_hash(text)
                        if h not in seen_hashes:
                            seen_hashes.add(h)
                            doc.add_paragraph(text)

            doc.save(file_path)
            print(f"Saved clean div content: {file_path}")

        except Exception as e:
            print(f"Error processing {url}: {e}")

    print(f"All files saved in folder: {output_folder}")

# Example usage:
if __name__ == "__main__":
    excel_file = "urls.xlsx"
    sheet_name = "Sheet1"
    url_column = "URL"
    output_folder = "div_content_pages"

    urls_from_excel_div_content(excel_file, sheet_name, url_column, output_folder)
