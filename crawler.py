
import pandas as pd
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

# Clean ads and irrelevant sections
def clean_article(article):
    for ad_tag in article.find_all(['div', 'section', 'aside'], class_=lambda x: x and ('ad' in x or 'advertisement' in x or 'promo' in x)):
        ad_tag.decompose()
    for ad_tag in article.find_all(['div', 'section', 'aside'], id=lambda x: x and ('ad' in x or 'advertisement' in x or 'promo' in x)):
        ad_tag.decompose()
    return article

# Sanitize file names
def sanitize_filename(name):
    name = re.sub(r'[^a-zA-Z0-9_-]', '_', name)
    return name[:50]

# Read URLs from Excel file
def read_urls_from_excel(file_path):
    df = pd.read_excel(file_path, engine='openpyxl')
    if 'URL' not in df.columns:
        raise ValueError("Excel file must contain a column named 'URL'")
    return df['URL'].dropna().tolist()

# Crawl URLs and save each article as a separate Word file
def crawl_and_save_separately(urls, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for url in urls:
        try:
            response = requests.get(url)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')

            article = soup.find('article') or soup.find('div', class_='content') or soup.find('div', class_='post')
            if not article:
                print(f"Could not find article content for {url}")
                continue

            article = clean_article(article)
            paragraphs = [p.get_text(strip=True) for p in article.find_all('p') if p.get_text(strip=True)]

            title = soup.title.string if soup.title else url
            filename = sanitize_filename(title) + ".docx"
            filepath = os.path.join(output_dir, filename)

            doc = Document()
            title_paragraph = doc.add_paragraph(title, style='Title')
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for para in paragraphs:
                p = doc.add_paragraph(para, style='Normal')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.runs[0]
                run.font.size = Pt(12)

            doc.save(filepath)
            print(f"Saved article from {url} to {filepath}")

        except Exception as e:
            print(f"Error processing {url}: {e}")

# Example usage:
excel_file = "urls.xlsx"  # Your Excel file with a column named 'URL'
output_dir = "articles"
urls = read_urls_from_excel(excel_file)
crawl_and_save_separately(urls, output_dir)
